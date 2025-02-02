import streamlit as st
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
import heapq
import io
from fpdf import FPDF
import docx

# Ensure necessary NLTK resources are available
def download_nltk_resources():
    nltk.download('punkt')
    nltk.download('stopwords')

def nltk_summarizer(raw_text):
    stopWords = set(stopwords.words("english"))
    word_frequencies = {}
    for word in nltk.word_tokenize(raw_text):
        if word.lower() not in stopWords:
            if word not in word_frequencies:
                word_frequencies[word] = 1
            else:
                word_frequencies[word] += 1

    maximum_frequency = max(word_frequencies.values())

    for word in word_frequencies:
        word_frequencies[word] = word_frequencies[word] / maximum_frequency

    sentence_list = nltk.sent_tokenize(raw_text)
    sentence_scores = {}
    for sent in sentence_list:
        for word in nltk.word_tokenize(sent.lower()):
            if word in word_frequencies:
                if len(sent.split(' ')) < 30:
                    if sent not in sentence_scores:
                        sentence_scores[sent] = word_frequencies[word]
                    else:
                        sentence_scores[sent] += word_frequencies[word]

    summary_sentences = heapq.nlargest(7, sentence_scores, key=sentence_scores.get)
    summary = ' '.join(summary_sentences)
    return summary, len(raw_text.split()), len(summary.split())

def generate_pdf(summary_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, summary_text)
    pdf_output = io.BytesIO()
    pdf.output(pdf_output, 'F')
    pdf_output.seek(0)
    return pdf_output

def generate_docx(summary_text):
    doc = docx.Document()
    doc.add_paragraph(summary_text)
    doc_output = io.BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)
    return doc_output

def process_text_input():
    st.subheader("Text Summarizer")
    raw_text = st.text_area("Enter Text Here")
    format_option = st.selectbox("Select Download Format", ["Text", "PDF", "DOCX"])
    if st.button("Summarize Text"):
        if raw_text.strip():
            with st.spinner("Summarizing..."):
                summary_result, input_length, summary_length = nltk_summarizer(raw_text)
            st.write("Summary:")
            st.write(summary_result)
            st.write(f"Input Length: {input_length} words")
            st.write(f"Summary Length: {summary_length} words")

            if format_option == "Text":
                st.download_button(label="Download Summary", data=summary_result, file_name="summary.txt", mime="text/plain")
            elif format_option == "PDF":
                pdf_data = generate_pdf(summary_result)
                st.download_button(label="Download Summary as PDF", data=pdf_data, file_name="summary.pdf", mime="application/pdf")
            elif format_option == "DOCX":
                docx_data = generate_docx(summary_result)
                st.download_button(label="Download Summary as DOCX", data=docx_data, file_name="summary.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.write("Please enter some text to summarize")

def process_file_upload():
    st.subheader("Document Summarizer")
    uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx"])
    format_option = st.selectbox("Select Download Format", ["Text", "PDF", "DOCX"])
    if st.button("Summarize Document"):
        raw_text = ""
        if uploaded_file is not None:
            try:
                if uploaded_file.type == "text/plain":
                    raw_text = uploaded_file.read().decode("utf-8")
                elif uploaded_file.type == "application/pdf":
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(uploaded_file)
                    for page in pdf_reader.pages:
                        raw_text += page.extract_text()
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc = docx.Document(uploaded_file)
                    raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                if raw_text.strip():
                    with st.spinner("Summarizing..."):
                        summary_result, input_length, summary_length = nltk_summarizer(raw_text)
                    st.write("Summary:")
                    st.write(summary_result)
                    st.write(f"Input Length: {input_length} words")
                    st.write(f"Summary Length: {summary_length} words")

                    if format_option == "Text":
                        st.download_button(label="Download Summary", data=summary_result, file_name="summary.txt", mime="text/plain")
                    elif format_option == "PDF":
                        pdf_data = generate_pdf(summary_result)
                        st.download_button(label="Download Summary as PDF", data=pdf_data, file_name="summary.pdf", mime="application/pdf")
                    elif format_option == "DOCX":
                        docx_data = generate_docx(summary_result)
                        st.download_button(label="Download Summary as DOCX", data=docx_data, file_name="summary.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.write("The uploaded document is empty")
            except Exception as e:
                st.write(f"Error processing file: {e}")
        else:
            st.write("Please upload a document to summarize")

def app_overview():
    st.title("Text Summarizer App")
    st.write("""
    Welcome to the Text Summarizer App!

    This app allows you to summarize text and documents with ease. 
    You can either enter text directly or upload a document to get a concise summary. 

    ### Features:
    - Summarize text by typing or pasting it into the text area.
    - Upload documents (TXT, PDF, DOCX) for summarization.
    - Download the summary in various formats (Text, PDF, DOCX).
    
    Use the sidebar to navigate between the text summarizer and document summarizer options.

    ### Instructions:
    1. Select the input type from the sidebar (Text Summarizer or Document Summarizer).
    2. Enter the text or upload a document.
    3. Click on "Summarize" to get the summary.
    4. Download the summary in your preferred format.
    """)

    if st.button("Let's Summarize"):
        st.session_state.show_summarizer = True

def main():
    download_nltk_resources()
    
    if 'show_summarizer' not in st.session_state:
        st.session_state.show_summarizer = False

    if not st.session_state.show_summarizer:
        app_overview()
    else:
        st.sidebar.title("Options")
        option = st.sidebar.selectbox(
            "Select Input Type",
            ("Text Summarizer", "Document Summarizer")
        )

        if option == "Text Summarizer":
            process_text_input()
        else:
            process_file_upload()

if __name__ == "__main__":
    main()
